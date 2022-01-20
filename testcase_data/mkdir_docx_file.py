# -*- encoding: utf-8 -*-
'''
@时间 ： 2022/1/9 18:32
@作者 ： 王齐涛
@文件名称： mkdir_docx_file.py 
'''
from docx import Document
from faker import Faker

fake = Faker(locale="zh_CN")    # 调用python的一个库来生成对应的测试数据
__txt_data = fake.texts(nb_texts=3, max_nb_chars=200, ext_word_list=None)  # 生成随机的文章，循环一次是2kb
__filename = fake.random_int()    # 随机数字，默认0~9999，可以通过设置min,max来设置
__folder_name = fake.name()     # 随机生成全名



document = Document()   # 新建文档
document.add_heading("标题内容，测试")
document.add_picture(r"G:\locust_data\s.jpg")
filename = str(fake.sha1())
filepath = f"M:\\AOMEIYUNDATA\\docx\\"+filename + ".docx"  # 文件生成的路径
for j in range(1, 20):
    document.add_paragraph(f"{__txt_data}+{j}")    # 添加段落内容
    document.save(f"{filepath}")
    print("生成文件成功")


